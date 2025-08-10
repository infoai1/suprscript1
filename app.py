import streamlit as st
from docx import Document
from io import BytesIO
from docx.shared import RGBColor

def detect_superscripts_and_find_refs(docx_bytes):
    document = Document(BytesIO(docx_bytes))
    superscripts = []

    # Detect superscript numbers
    for para in document.paragraphs:
        for run in para.runs:
            if run.font.superscript and run.text.strip().isdigit():
                superscripts.append((run.text.strip(), para.text.strip()))

    # Find paragraphs starting exactly with number + dot
    superscript_refs = []
    for num, sup_para_text in superscripts:
        options = [para.text.strip() for para in document.paragraphs if para.text.strip().startswith(f"{num}.")]
        superscript_refs.append((num, sup_para_text, options))

    return superscript_refs

def replace_superscript_with_reference(docx_bytes, selections):
    document = Document(BytesIO(docx_bytes))

    for para in document.paragraphs:
        for run in para.runs:
            if run.font.superscript and run.text.strip().isdigit():
                num = run.text.strip()
                if num in selections and selections[num]:
                    new_text = f"(Ref: {selections[num]})"
                    run.text = new_text
                    run.font.superscript = False
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red color

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

st.title("Superscript Reference Selector and Replacer")

uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])

if uploaded_file is not None:
    doc_bytes = uploaded_file.read()
    superscripts_with_options = detect_superscripts_and_find_refs(doc_bytes)

    st.write("Select the correct reference paragraph for each superscript below:")

    selections = {}
    for idx, (num, sup_text, options) in enumerate(superscripts_with_options):
        st.markdown(f"**Superscript {num} found in:**")
        st.write(f"> {sup_text}")

        if not options:
            st.write("No matching reference paragraphs found for this superscript. It will be left unchanged.")
            selections[num] = None
        elif len(options) == 1:
            st.write(f"Only one option found (auto-selected): {options[0]}")
            selections[num] = options
        else:
            choice = st.selectbox(f"Choose a reference for superscript {num}:", options, key=f"ref_{idx}")
            selections[num] = choice
        st.markdown("---")

    if st.button("Replace Superscripts and Download DOCX"):
        replaced_buffer = replace_superscript_with_reference(doc_bytes, selections)
        st.download_button(label="Download Modified DOCX",
                           data=replaced_buffer,
                           file_name="modified_references.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
